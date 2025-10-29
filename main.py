from docx import Document
import pypandoc
import pandas as pd
from datetime import datetime, timedelta
import time
import base64
import os
from PyPDF2 import PdfReader, PdfWriter, PdfMerger
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import red
import pytz

database = pd.read_csv('/content/Approved Appraisers.csv')

# This needs to figure out how to source the inputs while minimizing the amount of necessary inputs from user

def filter_by_type(database, type_value): # This filters our database based on the type of vendor we're looking for
    return database[database['Type'].str.contains(type_value, case=False, na=False)]

# This is our function for autofilling information from our csv of appraiser information

def autofill(data, type_value):
    database = filter_by_type(data, type_value)

    first = input('What is the first name of the vendor? If they are not on our list, type NA: ')

    if first == 'NA':
      first = input('What is the first name of this vendor? ')
      last = input('What is the last name of this vendor? ')
      email = input("What is the vendor's email? ")
      company = input('What company is the vendor with? ')
      inputs = {'First': first, 'Last': last, 'Email': email, 'Company': company}
      return inputs
    else:
      filtered = database[database['First'].str.contains(first, case=False, na=False)]

      if len(filtered) == 1:  # Single match
          inputs = filtered.iloc[0].to_dict()  # Convert the single row to a dictionary
      elif len(filtered) > 1:  # Multiple matches, ask for last name
          last = input('What is the last name of the vendor? ')
          filtered = filtered[filtered['Last'].str.contains(last, case=False, na=False)]
          if len(filtered) == 1:  # Single match now
              inputs = filtered.iloc[0].to_dict()
          else:
              return f"Number of cases found: {len(filtered)}"
      else:
          return f'Nobody named {first} found'

      return inputs  # Output is in a dict format

# Date builder is a function which will return the proper date in the future given some input and the duration requested (bds, weeks, or days). In current form, it doesn't adjust for holidays

def date_builder(days=0, weeks=0, business_days=0):
    tz = pytz.timezone('America/Los_Angeles')
    start_date = datetime.now(tz)
    if business_days > 0:
        current_date = start_date
        while business_days > 0:
            current_date += timedelta(days=1)
            # Skip weekends (Saturday=5, Sunday=6)
            if current_date.weekday() < 5:
                business_days -= 1
        # Keep as datetime
        delivery_date = current_date
    elif weeks > 0:
        delivery_date = datetime.today() + timedelta(weeks=weeks)
    else:
        delivery_date = datetime.today() + timedelta(days=days)

    dates = {
        'current_date': start_date.strftime('%-m/%-d/%Y'),
        'delivery_date': delivery_date.strftime('%-m/%-d/%Y')
    }
    return dates


# This allows a single call to orchestrate calls to the date builder and autofill so the end user can have an easy workflow

def input_builder(loan_type, letter_type):

  # Obtains info using autofill function to get info about vendor
  vendor_info = autofill(database, letter_type)

  print('Vendor Information', vendor_info)

  if letter_type != 'Sec':
    date_step = input('In how long does this need to be done? If business days, say "10 bds", if weeks, say "2 wks", if regular days, say "10 days". ')

    # Parse date input
    quantity = date_step.split(' ')[0]
    unit = date_step.split(' ')[1]

    if unit == 'wks':
      dates = date_builder(days=0, weeks=int(quantity), business_days=0)
      print(dates)
    elif unit == 'bds':
      dates = date_builder(days=0, weeks=0, business_days=int(quantity))
      print(dates)
    else:
      dates = date_builder(days=int(quantity), weeks=0, business_days=0)
      print(dates)

  if loan_type == '504': # If it's a 504 loan, we want to get the cdc company
    cdc_company = input("Who is the CDC Company? ")
    today_date = dates['current_date']
    loan_name = input('What is the loan name? ')
    loan_number = input('What is the loan number? ')
    property_address = input('What is the property address?')
    sqft = input('How many square feet for the property? ')
    fee = input('What is the fee being charged? ')
    delivery_date = dates['delivery_date']
    property_type = input('What is the property type? ')
    property_contact_name = input('Who is the property contact? ')
    property_contact_phone = input('What is the number of the property contact? ')
    item_to_send = 'TBD'
  else: # Not a 504, so we don't need special information, check for secondary appraisal
    if letter_type == 'Sec': # Condition for a secondary appraisal where we only need limited info
      tz = pytz.timezone('America/Los_Angeles')
      today_date = datetime.now(tz).strftime('%-m/%-d/%Y')
      print(today_date)
      loan_name = input('What is the loan name? ')
      loan_number = input('What is the loan number? ')
      property_address = input('What is the property address?')
      sqft = input('How many square feet for the property? ')
      fee = '600'
      delivery_date = '2 to 4 days'
      property_type = input('What is the property type? ')
      property_contact_name = 'NA'
      property_contact_phone = 'NA'
      item_to_send = 'TBD'
      cdc_company = 'N/A'
    else:
      today_date = dates['current_date']
      loan_name = input('What is the loan name? ')
      loan_number = input('What is the loan number? ')
      property_address = input('What is the property address?')
      sqft = input('How many square feet for the property? ')
      fee = input('What is the fee being charged? ')
      delivery_date = dates['delivery_date']
      property_type = input('What is the property type? ')
      property_contact_name = input('Who is the property contact? ')
      property_contact_phone = input('What is the number of the property contact? ')
      item_to_send = 'TBD'
      cdc_company = 'N/A'

  inputs = {
      "{{date}}": str(today_date),
      "{{contact_first_name}}": str(vendor_info['First']),
      "{{contact_last_name}}": str(vendor_info['Last']),
      "{{loan_name}}": loan_name,
      "{{company_name}}": str(vendor_info['Company']),
      "{{contact_email}}": str(vendor_info['Email']),
      "{{loan_number}}": loan_number,
      "{{property_address}}": property_address,
      "{{property_type}}": property_type,
      "{{sqft}}": sqft,
      "{{fee}}": fee,
      "{{delivery_date}}": str(delivery_date),
      "{{property_contact_name}}": property_contact_name,
      "{{property_contact_phone}}": property_contact_phone,
      "{{item_to_send}}": item_to_send,
      "{{cdc_company}}": cdc_company}

  return inputs

def engagement_builder():

  loan_type = input('What type of loan is this? Input either 7a, CC, or 504: ')
  letter_type = input('What kind of engagement letter is this? Choose App, Env, Sec, or SFR')

  inputs = input_builder(loan_type, letter_type)

  print(inputs)

  # Load proper document

  if letter_type == 'App':
    doc = Document(f'/content/{loan_type} - Appraisal Engagement Letter.docx') # Build doc object for editing
  elif letter_type == 'Sec':
    doc = Document(f'/content/{loan_type} - Appraisal Review Engagement Letter.docx') # Build doc object for editing
  else:
    doc = Document(f'/content/{loan_type} - Environmental Engagement Letter.docx') # Build doc object for editing

  # Current the replacements should be

  for paragraph in doc.paragraphs:
      for placeholder, replacement in inputs.items():
          if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement)



  header = doc.sections[0].header

  # Adjusts our header
  for paragraph in header.paragraphs:
      for placeholder, replacement in inputs.items():
          if placeholder in paragraph.text:
              paragraph.text = paragraph.text.replace(placeholder, replacement)

  # Save the updated document

  if letter_type == 'App' or letter_type =='SFR':
    doc.save(f"{inputs['{{loan_name}}']} Appraisal Engagement Letter.docx")
    print(f"Saved {inputs['{{loan_name}}']} Appraisal engagement letter!")
  elif letter_type == 'Sec':
    doc.save(f"{inputs['{{loan_name}}']} Appraisal Review  Engagement Letter.docx")
    print(f"Saved {inputs['{{loan_name}}']} Appraisal review engagement letter!")
  else:
    doc.save(f"{inputs['{{loan_name}}']} Phase 1 Engagement Letter.docx")
    print(f"Saved {inputs['{{loan_name}}']} Phase 1 engagement letter!")


# THIS FUNCTION IS FOR PRODUCING APPRAISAL AND PHASE 1 AT THE SAME TIME

# This allows a single call to orchestrate calls to the date builder and autofill so the end user can have an easy workflow
# This function is specifically for when I'm building two documents simultaneously.

def filter_by_type(database, type_value): # This filters our database based on the type of vendor we're looking for
    return database[database['Type'].str.contains(type_value, case=False, na=False)]

# This is our function for autofilling information from our csv of appraiser information

def autofill(data, type_value):
    database = filter_by_type(data, type_value)

    first = input('What is the first name of the vendor? If not in database, please type "NA": ')

    if first == 'NA':
        company = input('Please input company here: ')
        first = input('Please type first name here: ')
        last = input('Please input last name here: ')
        type_input = type_value
        email = input('Please input email here: ')
        region = 'NA'
        inputs = {'Company': company, 'First': first, 'Last': last, 'Type': type_value, 'Email': email, 'Region': region}
        return inputs

    filtered = database[database['First'].str.contains(first, case=False, na=False)]

    if len(filtered) == 1:  # Single match
        inputs = filtered.iloc[0].to_dict()  # Convert the single row to a dictionary
    elif len(filtered) > 1:  # Multiple matches, ask for last name
        last = input('What is the last name of the vendor? ')
        filtered = filtered[filtered['Last'].str.contains(last, case=False, na=False)]
        if len(filtered) == 1:  # Single match now
            inputs = filtered.iloc[0].to_dict()
        else: # This won't happen but if so, we can have an option to choose between options
            return f"Number of cases found: {len(filtered)}"
    else:
        return f'Nobody named {first} found'

    return inputs  # Output is in a dict format

# Date builder is a function which will return the proper date in the future given some input and the duration requested (bds, weeks, or days). In current form, it doesn't adjust for holidays

def date_builder(days=0, weeks=0, business_days=0):
    tz = pytz.timezone('America/Los_Angeles')  # Replace with your desired timezone
    start_date = datetime.now(tz)
    if business_days > 0:
        current_date = start_date
        while business_days > 0:
            current_date += timedelta(days=1)
            # Skip weekends (Saturday=5, Sunday=6)
            if current_date.weekday() < 5:
                business_days -= 1
        # Keep as datetime
        delivery_date = current_date
    elif weeks > 0:
        delivery_date = datetime.today() + timedelta(weeks=weeks)
    else:
        delivery_date = datetime.today() + timedelta(days=days)

    dates = {
        'current_date': start_date.strftime('%-m/%-d/%Y'),
        'delivery_date': delivery_date.strftime('%-m/%-d/%Y')
    }
    return dates

def dual_input_builder(loan_type):

  # Obtains info using autofill function to get info about the appraiser

  print('Information for Appraiser')

  app_info = autofill(database, type_value='App')
  print(app_info)

  # Obtain info using autofill function to get info about the environmental

  print('Information for Environmental')

  env_info = autofill(database, type_value='Env')
  print(env_info)

  # All date steps are conducted here, we build a function so we can streamline applying it to both appraisal and phase 1

  def dual_date_builder(date_step):

    # Parse date input
    quantity = date_step.split(' ')[0]
    unit = date_step.split(' ')[1]

    if unit == 'wks':
      dates = date_builder(days=0, weeks=int(quantity), business_days=0)
      print(dates)
    elif unit == 'bds':
      dates = date_builder(days=0, weeks=0, business_days=int(quantity))
      print(dates)
    else:
      dates = date_builder(days=int(quantity), weeks=0, business_days=0)
    return dates

  # Build the dates for the appraisal here
  app_date_step = input('In how long does the appraisal need to be done? If business days, say "10 bds", if weeks, say "2 wks", if regular days, say "10 days". ')
  app_dates = dual_date_builder(app_date_step)
  print(f"Appraisal Dates:\n Current Day: {app_dates['current_date']}\n Delivery Date: {app_dates['delivery_date']}")

  # Build the dates for the environmental here
  env_date_step = input('In how long does the phase 1 need to be done? If business days, say "10 bds", if weeks, say "2 wks", if regular days, say "10 days". ')
  env_dates = dual_date_builder(env_date_step)
  print(f"Environmental Dates:\n Current Day: {env_dates['current_date']}\n Delivery Date: {env_dates['delivery_date']}")

  print(loan_type)

  # Here loan_type is same for both

  if loan_type == '504': # If it's a 504 loan, we want to get the cdc company
    cdc_company = input("Who is the CDC Company? ")
    today_date = app_dates['current_date']
    loan_name = input('What is the loan name? ')
    loan_number = input('What is the loan number? ')
    property_address = input('What is the property address?')
    sqft = input('How many square feet for the property? ')
    app_fee = input('What is the fee being charged for the appraisal? ')
    env_fee = input('What is the fee being charged for the phase 1? ')
    app_delivery_date = app_dates['delivery_date']
    env_delivery_date = env_dates['delivery_date']
    property_type = input('What is the property type? ')
    property_contact_name = input('Who is the property contact? ')
    property_contact_phone = input('What is the number of the property contact? ')
    item_to_send = 'TBD'
  else: # Not a 504, so we don't need special information, check for secondary appraisal
      today_date = app_dates['current_date']
      loan_name = input('What is the loan name? ')
      loan_number = input('What is the loan number? ')
      property_address = input('What is the property address?')
      sqft = input('How many square feet for the property? ')
      app_fee = input('What is the fee being charged for the appraisal? ')
      env_fee = input('What is the fee being charged for the phase 1? ')
      app_delivery_date = app_dates['delivery_date']
      env_delivery_date = env_dates['delivery_date']
      property_type = input('What is the property type? ')
      property_contact_name = input('Who is the property contact? ')
      property_contact_phone = input('What is the number of the property contact? ')
      item_to_send = 'TBD'
      cdc_company = 'N/A'

  app_inputs = {
        "{{date}}": str(today_date),
        "{{contact_first_name}}": str(app_info['First']),
        "{{contact_last_name}}": str(app_info['Last']),
        "{{loan_name}}": loan_name,
        "{{company_name}}": str(app_info['Company']),
        "{{contact_email}}": str(app_info['Email']),
        "{{loan_number}}": loan_number,
        "{{property_address}}": property_address,
        "{{property_type}}": property_type,
        "{{sqft}}": sqft,
        "{{fee}}": app_fee,
        "{{delivery_date}}": str(app_delivery_date),
        "{{property_contact_name}}": property_contact_name,
        "{{property_contact_phone}}": property_contact_phone,
        "{{item_to_send}}": item_to_send,
        "{{cdc_company}}": cdc_company}
  env_inputs = {
        "{{date}}": str(today_date),
        "{{contact_first_name}}": str(env_info['First']),
        "{{contact_last_name}}": str(env_info['Last']),
        "{{loan_name}}": loan_name,
        "{{company_name}}": str(env_info['Company']),
        "{{contact_email}}": str(env_info['Email']),
        "{{loan_number}}": loan_number,
        "{{property_address}}": property_address,
        "{{property_type}}": property_type,
        "{{sqft}}": sqft,
        "{{fee}}": env_fee,
        "{{delivery_date}}": str(env_delivery_date),
        "{{property_contact_name}}": property_contact_name,
        "{{property_contact_phone}}": property_contact_phone,
        "{{item_to_send}}": item_to_send,
        "{{cdc_company}}": cdc_company}

  inputs = {
      'app_inputs': app_inputs,
      'env_inputs': env_inputs}

  return inputs

def dual_engagement_builder():

  # Loan type

  loan_type = input('What type of loan is this? Input either 7a, CC, or 504: ')

  # This will return inputs for both our appraisal and phase 1
  inputs = dual_input_builder(loan_type)

  # Build a basic function for carrying out this process twice
  def dual_engagement_helper(inputs, letter_type):
    if letter_type == 'App':
      doc = Document(f'/content/{loan_type} - Appraisal Engagement Letter.docx') # Build doc object for editing
    else:
      doc = Document(f'/content/{loan_type} - Environmental Engagement Letter.docx') # Build doc object for editing

    for paragraph in doc.paragraphs:
        for placeholder, replacement in inputs.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)

    header = doc.sections[0].header

    # Adjusts our header
    for paragraph in header.paragraphs:
        for placeholder, replacement in inputs.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)

    # Save the updated document

    if letter_type == 'App' or letter_type =='SFR':
      doc.save(f"{inputs['{{loan_name}}']} Appraisal Engagement Letter.docx")
      print(f"Saved {inputs['{{loan_name}}']} Appraisal engagement letter!")
    else:
      doc.save(f"{inputs['{{loan_name}}']} Phase 1 Engagement Letter.docx")
      print(f"Saved {inputs['{{loan_name}}']} Phase 1 engagement letter!")

  dual_engagement_helper(inputs['app_inputs'], 'App')
  dual_engagement_helper(inputs['env_inputs'], 'Env')