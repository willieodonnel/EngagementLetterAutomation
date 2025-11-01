"""
Document Generation Module for Engagement Letter Automation

This module handles the generation of engagement letters from templates.
It takes processed data in JSON format and creates customized Word documents
by replacing placeholders in templates with actual values.
"""

import os
from typing import Dict, Any, List, Tuple
from docx import Document


def generate_engagement_letter(
    data: Dict[str, Any],
    template_dir: str = "Templates",
    output_dir: str = ".",
    output_filename: str = None
) -> str:
    """
    Generate a single engagement letter from data and template.

    Args:
        data: Dictionary containing all engagement data
        template_dir: Directory containing Word templates
        output_dir: Directory to save generated documents
        output_filename: Optional custom output filename

    Returns:
        Path to the generated document
    """

    # Extract key information
    loan_type = data["loan_type"]
    letter_type = data["letter_type"]
    loan_name = data["loan"]["loan_name"]

    # Determine template path
    template_path = _get_template_path(loan_type, letter_type, template_dir)

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    # Convert data to placeholder format
    placeholders = _data_to_placeholders(data)

    # Load and process the template
    doc = Document(template_path)
    _replace_placeholders_in_document(doc, placeholders)

    # Generate output filename if not provided
    if not output_filename:
        letter_type_display = _get_letter_type_display(letter_type)
        output_filename = f"{loan_name} {letter_type_display} Engagement Letter.docx"

    # Save the document
    output_path = os.path.join(output_dir, output_filename)
    doc.save(output_path)

    print(f"Generated: {output_path}")
    return output_path


def generate_dual_engagement_letters(
    dual_data: Dict[str, Any],
    template_dir: str = "Templates",
    output_dir: str = "."
) -> Tuple[str, str]:
    """
    Generate both appraisal and environmental engagement letters.

    Args:
        dual_data: Dictionary containing data for both letters
        template_dir: Directory containing Word templates
        output_dir: Directory to save generated documents

    Returns:
        Tuple of paths to generated documents (appraisal_path, environmental_path)
    """

    loan_type = dual_data["loan_type"]
    loan_name = dual_data["shared"]["loan"]["loan_name"]

    # Generate appraisal letter
    app_data = {
        "loan_type": loan_type,
        "letter_type": "APP",
        "vendor": dual_data["appraisal"]["vendor"],
        "dates": dual_data["appraisal"]["dates"],
        "loan": dual_data["shared"]["loan"],
        "property": dual_data["shared"]["property"]
    }

    # Add appraisal fee to property data if it exists
    if "fee" in dual_data["appraisal"]["dates"]:
        app_data["property"]["fee"] = dual_data["appraisal"]["dates"]["fee"]

    app_path = generate_engagement_letter(
        app_data,
        template_dir,
        output_dir,
        f"{loan_name} Appraisal Engagement Letter.docx"
    )

    # Generate environmental letter
    env_data = {
        "loan_type": loan_type,
        "letter_type": "ENV",
        "vendor": dual_data["environmental"]["vendor"],
        "dates": dual_data["environmental"]["dates"],
        "loan": dual_data["shared"]["loan"],
        "property": dual_data["shared"]["property"]
    }

    # Add environmental fee to property data if it exists
    if "fee" in dual_data["environmental"]["dates"]:
        env_data["property"]["fee"] = dual_data["environmental"]["dates"]["fee"]

    env_path = generate_engagement_letter(
        env_data,
        template_dir,
        output_dir,
        f"{loan_name} Phase 1 Engagement Letter.docx"
    )

    return app_path, env_path


def _get_template_path(loan_type: str, letter_type: str, template_dir: str) -> str:
    """
    Determine the correct template path based on loan and letter types.

    Args:
        loan_type: Type of loan ('7a', '504', 'CC')
        letter_type: Type of letter
        template_dir: Base directory for templates

    Returns:
        Full path to the template file
    """

    loan_type = loan_type.upper()
    letter_type = letter_type.upper()

    # Map letter types to template names
    if letter_type in ['APP', 'APPRAISAL']:
        template_name = f"{loan_type} - Appraisal Engagement Letter.docx"
    elif letter_type in ['SEC', 'SECONDARY']:
        template_name = f"{loan_type} - Appraisal Review Engagement Letter.docx"
    elif letter_type in ['ENV', 'ENVIRONMENTAL']:
        template_name = f"{loan_type} - Environmental Engagement Letter.docx"
    elif letter_type == 'PHASE 1':
        template_name = f"{loan_type} - Phase 1 Engagement Letter.docx"
    elif letter_type == 'PHASE 2':
        template_name = f"{loan_type} - Phase 2 Engagement Letter.docx"
    elif letter_type == 'SFR':
        # Single Family Residence - uses appraisal template
        template_name = f"{loan_type} - Appraisal Engagement Letter.docx"
    else:
        # Default to environmental template for unknown types
        template_name = f"{loan_type} - Environmental Engagement Letter.docx"

    return os.path.join(template_dir, template_name)


def _data_to_placeholders(data: Dict[str, Any]) -> Dict[str, str]:
    """
    Convert structured data to placeholder key-value pairs.

    Args:
        data: Structured engagement data

    Returns:
        Dictionary mapping placeholder strings to replacement values
    """

    placeholders = {}

    # Add vendor placeholders
    if "vendor" in data:
        vendor = data["vendor"]
        placeholders["{{contact_first_name}}"] = vendor.get("first_name", "")
        placeholders["{{contact_last_name}}"] = vendor.get("last_name", "")
        placeholders["{{company_name}}"] = vendor.get("company", "")
        placeholders["{{contact_email}}"] = vendor.get("email", "")

    # Add date placeholders
    if "dates" in data:
        dates = data["dates"]
        placeholders["{{date}}"] = dates.get("current_date", "")
        placeholders["{{delivery_date}}"] = dates.get("delivery_date", "")

    # Add loan placeholders
    if "loan" in data:
        loan = data["loan"]
        placeholders["{{loan_name}}"] = loan.get("loan_name", "")
        placeholders["{{loan_number}}"] = loan.get("loan_number", "")
        placeholders["{{cdc_company}}"] = loan.get("cdc_company", "N/A")

    # Add property placeholders
    if "property" in data:
        property_data = data["property"]
        placeholders["{{property_address}}"] = property_data.get("property_address", "")
        placeholders["{{property_type}}"] = property_data.get("property_type", "")
        placeholders["{{sqft}}"] = property_data.get("sqft", "")
        placeholders["{{fee}}"] = property_data.get("fee", "")
        placeholders["{{property_contact_name}}"] = property_data.get("property_contact_name", "")
        placeholders["{{property_contact_phone}}"] = property_data.get("property_contact_phone", "")
        placeholders["{{item_to_send}}"] = property_data.get("item_to_send", "TBD")

    return placeholders


def _replace_placeholders_in_document(doc: Document, placeholders: Dict[str, str]) -> None:
    """
    Replace all placeholders in a Word document with actual values.

    Args:
        doc: Word document object
        placeholders: Dictionary mapping placeholder strings to replacement values
    """

    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, replacement in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)

    # Replace in headers
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            for placeholder, replacement in placeholders.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, replacement)

        # Also check first page header if different
        if hasattr(section, 'first_page_header'):
            for paragraph in section.first_page_header.paragraphs:
                for placeholder, replacement in placeholders.items():
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, replacement)

    # Replace in footers
    for section in doc.sections:
        footer = section.footer
        for paragraph in footer.paragraphs:
            for placeholder, replacement in placeholders.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, replacement)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, replacement in placeholders.items():
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, replacement)


def _get_letter_type_display(letter_type: str) -> str:
    """
    Get the display name for a letter type.

    Args:
        letter_type: Internal letter type code

    Returns:
        Human-readable letter type name
    """

    letter_type = letter_type.upper()

    display_names = {
        "APP": "Appraisal",
        "APPRAISAL": "Appraisal",
        "SEC": "Appraisal Review",
        "SECONDARY": "Appraisal Review",
        "ENV": "Environmental",
        "ENVIRONMENTAL": "Environmental",
        "PHASE 1": "Phase 1",
        "PHASE 2": "Phase 2",
        "SFR": "Single Family Residence"
    }

    return display_names.get(letter_type, letter_type)


def validate_template_exists(loan_type: str, letter_type: str, template_dir: str = "Templates") -> bool:
    """
    Check if a template exists for the given loan and letter types.

    Args:
        loan_type: Type of loan
        letter_type: Type of letter
        template_dir: Directory containing templates

    Returns:
        True if template exists, False otherwise
    """

    template_path = _get_template_path(loan_type, letter_type, template_dir)
    return os.path.exists(template_path)


def list_available_templates(template_dir: str = "Templates") -> List[str]:
    """
    List all available template files.

    Args:
        template_dir: Directory containing templates

    Returns:
        List of template filenames
    """

    if not os.path.exists(template_dir):
        return []

    templates = []
    for filename in os.listdir(template_dir):
        if filename.endswith('.docx') and not filename.startswith('~'):
            templates.append(filename)

    return sorted(templates)


def generate_from_json(
    json_path: str,
    template_dir: str = "Templates",
    output_dir: str = "."
) -> str:
    """
    Generate an engagement letter from a JSON file.

    Args:
        json_path: Path to JSON file containing engagement data
        template_dir: Directory containing Word templates
        output_dir: Directory to save generated documents

    Returns:
        Path to the generated document
    """

    import json

    with open(json_path, 'r') as f:
        data = json.load(f)

    # Check if it's dual data
    if "appraisal" in data and "environmental" in data:
        paths = generate_dual_engagement_letters(data, template_dir, output_dir)
        return f"Generated: {paths[0]} and {paths[1]}"
    else:
        return generate_engagement_letter(data, template_dir, output_dir)